from __future__ import annotations

import os
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber
from docx import Document as DocxDocument


__all__ = ["parse_document", "postprocess_document", "parse_docx_document"]


TOC_RE_WITH_NUM = re.compile(
    r"^(?P<num>\d+(\.\d+)*)\s+(?P<title>.+?)\s+\.{3,}\s*(?P<page>\d+)\s*$"
)
TOC_RE_NO_NUM = re.compile(r"^(?P<title>.+?)\s+\.{3,}\s*(?P<page>\d+)\s*$")
HEADING_RE = re.compile(r"^(?P<num>\d+(\.\d+)*)(\.?)\s+(?P<title>.+)$")
PAGE_NUMBER_LINE_RE = re.compile(r"^\s*\d+\s*$")


@dataclass
class TocItem:
    raw_title: str
    number: Optional[str]
    level: int
    page: int


@dataclass
class PdfSectionFlat:
    raw_title: str
    number: Optional[str]
    level: int
    start_page: int
    end_page: int
    text: str


@dataclass
class HeadingItem:
    level: int
    title: str
    para_idx: int
    number: Optional[str]


@dataclass
class DocxSectionFlat:
    raw_title: str
    number: Optional[str]
    level: int
    text: str
    start_page: Optional[int]
    end_page: Optional[int]


def _normalize_id_from_title(title: str) -> str:
    t = title.lower().strip()
    t = re.sub(r"\s+", "_", t)
    t = re.sub(r"[^a-zа-я0-9_\.]+", "", t)
    return (t[:50] or "section").strip("_.")


def _clean_text_page_numbers(text: str) -> str:
    lines = text.splitlines()
    cleaned = [line for line in lines if not PAGE_NUMBER_LINE_RE.match(line)]
    return "\n".join(cleaned)


def _find_toc_page_index(pages: List[str]) -> Optional[int]:
    for i, text in enumerate(pages):
        if not text:
            continue
        if "СОДЕРЖАНИЕ" in text or "ОГЛАВЛЕНИЕ" in text:
            return i
    return None


def _extract_pdf_pages(path: str) -> List[str]:
    pages: List[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return pages


def _parse_toc_items(text: str) -> List[TocItem]:
    items: List[TocItem] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        m = TOC_RE_WITH_NUM.match(line)
        if m:
            num = m.group("num")
            title = m.group("title").strip()
            page = int(m.group("page"))
            level = 2 if "." in num else 1
            items.append(TocItem(raw_title=title, number=num, level=level, page=page))
            continue

        m2 = TOC_RE_NO_NUM.match(line)
        if m2:
            title = m2.group("title").strip()
            page = int(m2.group("page"))
            items.append(TocItem(raw_title=title, number=None, level=1, page=page))
            continue

    return items


def _split_pages_by_toc_items(pages: List[str], items: List[TocItem]) -> List[PdfSectionFlat]:
    sections: List[PdfSectionFlat] = []

    for i, item in enumerate(items):
        start_page = item.page
        start_idx = max(0, start_page - 1)

        if i + 1 < len(items):
            next_page = items[i + 1].page
            end_idx = max(start_idx, next_page - 2)
        else:
            end_idx = len(pages) - 1

        text = "\n".join(pages[start_idx : end_idx + 1])

        sections.append(
            PdfSectionFlat(
                raw_title=item.raw_title,
                number=item.number,
                level=item.level,
                start_page=start_idx + 1,
                end_page=end_idx + 1,
                text=text,
            )
        )

    return sections


def _build_hierarchical_sections_from_flat(
    sections_flat: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    result: List[Dict[str, Any]] = []
    current_parent: Optional[Dict[str, Any]] = None

    for sec in sections_flat:
        level = int(sec.get("level") or 1)
        number = sec.get("number")
        title = str(sec.get("title") or sec.get("raw_title") or "")
        sec_id = number or _normalize_id_from_title(title)

        if level == 1:
            node: Dict[str, Any] = {
                "id": sec_id,
                "number": number,
                "level": 1,
                "title": title,
                "start_page": sec.get("start_page"),
                "end_page": sec.get("end_page"),
                "text": sec.get("text", ""),
                "subsections": [],
            }
            result.append(node)
            current_parent = node
        else:
            node = {
                "id": sec_id,
                "number": number,
                "level": 2,
                "title": title,
                "start_page": sec.get("start_page"),
                "end_page": sec.get("end_page"),
                "text": sec.get("text", ""),
            }
            if current_parent is not None:
                current_parent["subsections"].append(node)
            else:
                result.append(node)

    return result

GROUP_RE = re.compile(r"[А-Яа-я]{2,5}-\d{3}")


def _guess_title_and_author_from_pdf_pages(pages: List[str]) -> Tuple[str, str]:
    title = "UNKNOWN_TITLE"
    author = "UNKNOWN_AUTHOR"

    if not pages or len(pages) < 2:
        return title, author

    for page in pages[1:]:
        if not page:
            continue

        for line in page.splitlines():
            line = line.strip()
            if not line:
                continue

            if "ВАРИАНТ" in line.upper():
                title = re.sub(r"(?i)вариант", "", line).strip(" :.-")
                break
        if title != "UNKNOWN_TITLE":
            break

    for page in pages:
        if not page:
            continue

        for line in page.splitlines():
            line = line.strip()
            if not line:
                continue

            if "СТУДЕНТ ГРУППЫ" in line.upper():
                rest = re.split(r"(?i)студент группы", line, maxsplit=1)[-1].strip()
                parts = rest.split()

                if parts and GROUP_RE.fullmatch(parts[0]):
                    author = " ".join(parts[1:]).strip()
                else:
                    author = rest

                if author:
                    return title, author

    return title, author

def _parse_pdf_document(path: str) -> Dict[str, Any]:
    pages = _extract_pdf_pages(path)

    toc_page_idx = _find_toc_page_index(pages)
    if toc_page_idx is None:
        raise ValueError("Не удалось найти страницу с содержанием или оглавлением")

    toc_items = _parse_toc_items(pages[toc_page_idx])
    if not toc_items:
        raise ValueError("Содержание найдено, но строки распарсить не удалось")

    flat = _split_pages_by_toc_items(pages, toc_items)
    flat_dicts: List[Dict[str, Any]] = [
        {
            "title": s.raw_title,
            "raw_title": s.raw_title,
            "number": s.number,
            "level": s.level,
            "start_page": s.start_page,
            "end_page": s.end_page,
            "text": s.text,
        }
        for s in flat
    ]
    sections = _build_hierarchical_sections_from_flat(flat_dicts)
    title, author = _guess_title_and_author_from_pdf_pages(pages)

    return {"title": title, "author": author, "type": "курсовой проект", "sections": sections}


def _extract_headings(doc: DocxDocument) -> List[HeadingItem]:
    headings: List[HeadingItem] = []

    for i, p in enumerate(doc.paragraphs):
        style_name = (p.style.name or "").lower()
        if not (style_name.startswith("heading") or "заголовок" in style_name):
            continue

        text = p.text.strip()
        if not text:
            continue

        m = HEADING_RE.match(text)
        if m:
            num = m.group("num")
            title = m.group("title").strip()
        else:
            num = None
            title = text

        if "heading 1" in style_name or "заголовок 1" in style_name:
            level = 1
        elif "heading 2" in style_name or "заголовок 2" in style_name:
            level = 2
        else:
            level = 2

        headings.append(HeadingItem(level=level, title=title, para_idx=i, number=num))

    return headings


def _assign_missing_subnumbers(headings: List[HeadingItem]) -> None:
    current_parent_num: Optional[str] = None
    sub_counter = 0

    for h in headings:
        if h.level == 1:
            current_parent_num = h.number
            sub_counter = 0
        elif h.level == 2:
            if h.number is None and current_parent_num is not None:
                sub_counter += 1
                h.number = f"{current_parent_num}.{sub_counter}"


def _build_flat_docx_sections(doc: DocxDocument, headings: List[HeadingItem]) -> List[DocxSectionFlat]:
    sections: List[DocxSectionFlat] = []

    for idx, h in enumerate(headings):
        start_idx = h.para_idx
        if idx + 1 < len(headings):
            end_idx = headings[idx + 1].para_idx - 1
        else:
            end_idx = len(doc.paragraphs) - 1

        text = "\n".join(p.text for p in doc.paragraphs[start_idx : end_idx + 1])
        text = _clean_text_page_numbers(text)

        sections.append(
            DocxSectionFlat(
                raw_title=h.title,
                number=h.number,
                level=h.level,
                text=text,
                start_page=None,
                end_page=None,
            )
        )

    return sections


def _guess_title_and_author_from_docx(doc: DocxDocument) -> Tuple[str, str]:
    title = "UNKNOWN_TITLE"
    author = "UNKNOWN_AUTHOR"

    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for t in texts:
        if "ВАРИАНТ" in t.upper():
            title = re.sub(r"(?i)вариант", "", t).strip(" :.-")
            break

    for t in texts:
        if "СТУДЕНТ ГРУППЫ" in t.upper():
            rest = re.split(r"(?i)студент группы", t, maxsplit=1)[-1].strip()
            parts = rest.split()

            if parts and GROUP_RE.fullmatch(parts[0]):
                author = " ".join(parts[1:]).strip()
            else:
                author = rest

            if author:
                break

    return title, author


def parse_docx_document(path: str) -> Dict[str, Any]:
    doc = DocxDocument(path)

    headings = _extract_headings(doc)
    if not headings:
        raise ValueError("В DOCX не найдены заголовки (Heading 1 или Heading 2).")

    _assign_missing_subnumbers(headings)
    flat = _build_flat_docx_sections(doc, headings)

    flat_dicts: List[Dict[str, Any]] = [
        {
            "title": s.raw_title,
            "raw_title": s.raw_title,
            "number": s.number,
            "level": s.level,
            "start_page": s.start_page,
            "end_page": s.end_page,
            "text": s.text,
        }
        for s in flat
    ]

    sections = _build_hierarchical_sections_from_flat(flat_dicts)
    title, author = _guess_title_and_author_from_docx(doc)

    return {"title": title, "author": author, "type": "курсовой проект", "sections": sections}


def parse_document(path: str) -> Dict[str, Any]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _parse_pdf_document(path)
    if ext == ".docx":
        return parse_docx_document(path)
    raise ValueError(f"Неподдерживаемое расширение: {ext}")


def _sync_section_pages(section: Dict[str, Any]) -> None:
    subsections: List[Dict[str, Any]] = section.get("subsections") or []
    if not subsections:
        return

    sub_start_pages = [s.get("start_page") for s in subsections if s.get("start_page") is not None]
    sub_end_pages = [s.get("end_page") for s in subsections if s.get("end_page") is not None]
    if not sub_start_pages or not sub_end_pages:
        return

    min_start = min(sub_start_pages)
    max_end = max(sub_end_pages)

    cur_start = section.get("start_page")
    cur_end = section.get("end_page")

    section["start_page"] = min_start if cur_start is None else min(cur_start, min_start)
    section["end_page"] = max_end if cur_end is None else max(cur_end, max_end)


def postprocess_document(doc: Dict[str, Any]) -> Dict[str, Any]:
    sections: List[Dict[str, Any]] = doc.get("sections", [])

    for section in sections:
        if isinstance(section.get("text"), str):
            section["text"] = _clean_text_page_numbers(section["text"])

        subsections: List[Dict[str, Any]] = section.get("subsections") or []
        for sub in subsections:
            if isinstance(sub.get("text"), str):
                sub["text"] = _clean_text_page_numbers(sub["text"])

        if section.get("level") == 1:
            _sync_section_pages(section)

    return doc
